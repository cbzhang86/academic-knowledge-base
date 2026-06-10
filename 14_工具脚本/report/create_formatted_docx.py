#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中文学术论文 .docx 格式转换器 (Skill 版)
=========================================
统一接口: python create_formatted_docx.py <action> [--input <input.md>] [--output <output.docx>] [--style <style>]

action:
  convert  - 将 .md 格式化为 .docx
  styles   - 列出所有可用格式模板

--style 可选: default, 学报, 北核, 学位论文, csse
"""

import re
import sys
import json
import argparse
from pathlib import Path

# ============================================================
# 格式配置库（方案 B - 可扩展）
# 每种格式定义：
#   page_margin: 页边距 cm [top, bottom, left, right]
#   body:       {font, size_pt, line_spacing, indent_pt}
#   heading:    {level: {font, size_pt, align, bold, space_before_pt, space_after_pt}}
#   quote:      {font, size_pt, left_indent_cm}
#   code:       {font, size_pt, line_spacing}
#   abstract:   {font, size_pt, label_bold}
# ============================================================

STYLES = {
    # --------------------------------------------------
    # 默认（通用学术论文）
    # --------------------------------------------------
    "default": {
        "label": "通用学术论文",
        "desc": "宋体正文/黑体标题，1.5倍行距，缩进2字符",
        "page_margin_cm": [2.54, 2.54, 3.17, 3.17],
        "body":       {"font": "宋体", "size_pt": 12, "line_spacing": 1.5, "indent_pt": 24},
        "bold":       {"font": "宋体", "size_pt": 12},
        "heading": {
            1: {"font": "黑体", "size_pt": 16, "align": "center", "bold": True, "space_before": 12, "space_after": 6},
            2: {"font": "黑体", "size_pt": 14, "align": "left",   "bold": True, "space_before": 12, "space_after": 6},
            3: {"font": "黑体", "size_pt": 12, "align": "left",   "bold": True, "space_before": 6,  "space_after": 3},
            4: {"font": "黑体", "size_pt": 12, "align": "left",   "bold": True, "space_before": 6,  "space_after": 3},
        },
        "bullet":     {"font": "宋体", "size_pt": 12, "line_spacing": 1.5},
        "quote":      {"font": "楷体", "size_pt": 11, "line_spacing": 1.25, "left_indent_cm": 1.0},
        "code":       {"font": "宋体", "size_pt": 9,  "line_spacing": 1.0},
        "first_line_indent": True,
    },

    # --------------------------------------------------
    # 国内大学学报（如深圳大学学报、西北大学学报等）
    # --------------------------------------------------
    "学报": {
        "label": "大学学报",
        "desc": "宋体正文/黑体标题，与《深圳大学学报》《西北大学学报》等格式接轨",
        "page_margin_cm": [2.54, 2.54, 3.17, 3.17],
        "body":       {"font": "宋体", "size_pt": 12, "line_spacing": 1.5, "indent_pt": 24},
        "bold":       {"font": "宋体", "size_pt": 12},
        "heading": {
            1: {"font": "黑体", "size_pt": 14, "align": "center", "bold": True, "space_before": 12, "space_after": 6},
            2: {"font": "黑体", "size_pt": 12, "align": "left",   "bold": True, "space_before": 10, "space_after": 4},
            3: {"font": "楷体", "size_pt": 12, "align": "left",   "bold": True, "space_before": 6,  "space_after": 3},
            4: {"font": "楷体", "size_pt": 12, "align": "left",   "bold": False,"space_before": 6,  "space_after": 3},
        },
        "bullet":     {"font": "宋体", "size_pt": 12, "line_spacing": 1.5},
        "quote":      {"font": "楷体", "size_pt": 11, "line_spacing": 1.25, "left_indent_cm": 1.0},
        "code":       {"font": "宋体", "size_pt": 9,  "line_spacing": 1.0},
        "first_line_indent": True,
    },

    # --------------------------------------------------
    # 北大中文核心期刊（常见格式）
    # --------------------------------------------------
    "北核": {
        "label": "北大核心期刊",
        "desc": "五号宋体/双栏，标题黑体，适用于多数北大核心期刊",
        "page_margin_cm": [2.0, 2.0, 2.0, 2.0],
        "body":       {"font": "宋体", "size_pt": 10.5, "line_spacing": 1.25, "indent_pt": 21},
        "bold":       {"font": "宋体", "size_pt": 10.5},
        "heading": {
            1: {"font": "黑体", "size_pt": 14, "align": "center", "bold": True, "space_before": 10, "space_after": 5},
            2: {"font": "黑体", "size_pt": 12, "align": "left",   "bold": True, "space_before": 8,  "space_after": 4},
            3: {"font": "黑体", "size_pt": 10.5, "align": "left", "bold": True, "space_before": 6,  "space_after": 2},
            4: {"font": "楷体", "size_pt": 10.5, "align": "left", "bold": True, "space_before": 6,  "space_after": 2},
        },
        "bullet":     {"font": "宋体", "size_pt": 10.5, "line_spacing": 1.25},
        "quote":      {"font": "楷体", "size_pt": 10, "line_spacing": 1.0, "left_indent_cm": 0.8},
        "code":       {"font": "宋体", "size_pt": 8,  "line_spacing": 1.0},
        "first_line_indent": True,
    },

    # --------------------------------------------------
    # 学位论文（硕博）
    # --------------------------------------------------
    "学位论文": {
        "label": "学位论文",
        "desc": "宋体正文小四/标题黑体三号，适合硕博学位论文",
        "page_margin_cm": [2.54, 2.54, 3.17, 3.17],
        "body":       {"font": "宋体", "size_pt": 12, "line_spacing": 1.5, "indent_pt": 24},
        "bold":       {"font": "宋体", "size_pt": 12},
        "heading": {
            1: {"font": "黑体", "size_pt": 18, "align": "center", "bold": True, "space_before": 18, "space_after": 12},
            2: {"font": "黑体", "size_pt": 16, "align": "left",   "bold": True, "space_before": 15, "space_after": 8},
            3: {"font": "黑体", "size_pt": 14, "align": "left",   "bold": True, "space_before": 12, "space_after": 6},
            4: {"font": "黑体", "size_pt": 12, "align": "left",   "bold": True, "space_before": 8,  "space_after": 4},
        },
        "bullet":     {"font": "宋体", "size_pt": 12, "line_spacing": 1.5},
        "quote":      {"font": "楷体", "size_pt": 11, "line_spacing": 1.25, "left_indent_cm": 1.2},
        "code":       {"font": "宋体", "size_pt": 9,  "line_spacing": 1.0},
        "first_line_indent": True,
    },

    # --------------------------------------------------
    # CSSCI 南大核心（常见格式）
    # --------------------------------------------------
    "csse": {
        "label": "CSSCI 南大核心",
        "desc": "五号宋体/双栏，标题黑体，适合 CSSCI 来源期刊",
        "page_margin_cm": [2.0, 2.0, 1.8, 1.8],
        "body":       {"font": "宋体", "size_pt": 10.5, "line_spacing": 1.25, "indent_pt": 21},
        "bold":       {"font": "宋体", "size_pt": 10.5},
        "heading": {
            1: {"font": "黑体", "size_pt": 14, "align": "center", "bold": True, "space_before": 10, "space_after": 5},
            2: {"font": "黑体", "size_pt": 12, "align": "left",   "bold": True, "space_before": 8,  "space_after": 4},
            3: {"font": "黑体", "size_pt": 10.5, "align": "left", "bold": True, "space_before": 6,  "space_after": 2},
            4: {"font": "楷体", "size_pt": 10.5, "align": "left", "bold": True, "space_before": 6,  "space_after": 2},
        },
        "bullet":     {"font": "宋体", "size_pt": 10.5, "line_spacing": 1.25},
        "quote":      {"font": "楷体", "size_pt": 10, "line_spacing": 1.0, "left_indent_cm": 0.8},
        "code":       {"font": "宋体", "size_pt": 8,  "line_spacing": 1.0},
        "first_line_indent": True,
    },
}


# ============================================================
# 核心功能
# ============================================================

def _align_from_str(s):
    """将 align 字符串转为 docx 枚举值"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(s, WD_ALIGN_PARAGRAPH.LEFT)


def _add_run(p, text, font_name, size_pt, bold=False, italic=False):
    """添加带中文字体设置的 run"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = p.add_run(text)
    run.font.name = font_name

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    return run


def _set_default_style(doc, cfg_body):
    """设置文档默认样式"""
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    style = doc.styles['Normal']
    font = style.font
    font.name = cfg_body["font"]
    font.size = Pt(cfg_body["size_pt"])

    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cfg_body["font"])

    pf = style.paragraph_format
    pf.line_spacing = cfg_body["line_spacing"]


def _parse_md(md_text):
    """将 Markdown 解析为结构化段落列表"""
    # 去除 YAML frontmatter
    if md_text.startswith('---'):
        end = md_text.find('---', 3)
        if end != -1:
            md_text = md_text[end + 3:].strip()

    lines = md_text.split('\n')
    paras = []

    # 表格收集：连续以 | 开头或包含 | 且含有 --- 分隔符的行
    table_lines = []

    def flush_table():
        if table_lines:
            # 从收集的行中提取表头和数据
            # 过滤掉分隔行（|---|）
            data_rows = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|?$', l)
                         or '---' not in l]
            # 如果只有分隔行没有数据，放弃
            if len(data_rows) >= 2:
                paras.append({'type': 'table', 'rows': data_rows})
            table_lines.clear()

    for line in lines:
        s = line.strip()

        # 检测表格行：以 | 开头
        if s.startswith('|'):
            table_lines.append(s)
            continue
        else:
            flush_table()

        if not s:
            paras.append({'type': 'spacer'})
        elif s.startswith('#'):
            level = len(s) - len(s.lstrip('#'))
            paras.append({'type': 'heading', 'level': min(level, 4), 'text': s.lstrip('#').strip()})
        elif s.startswith('---') or s.startswith('***'):
            paras.append({'type': 'hr'})
        elif s.startswith('- ') or s.startswith('* '):
            paras.append({'type': 'bullet', 'text': s[2:]})
        elif s.startswith('>'):
            paras.append({'type': 'quote', 'text': s.lstrip('>').strip()})
        elif s.startswith('```'):
            paras.append({'type': 'codeblock', 'text': s})
        else:
            paras.append({'type': 'paragraph', 'text': s})

    flush_table()  # 文件末尾的表格
    return paras


def md_to_docx(md_path, output_path, style_name="default"):
    """
    将 Markdown 论文格式化为符合中文学术规范的 .docx
    ---
    统一 Skill 接口，供 pipeline / CLI 调用。
    stdout: JSON {"status": "ok", "output": "...", "style": "...", "pages_approx": N}
    stderr: 日志
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    # 读取样式配置
    if style_name not in STYLES:
        available = ", ".join(STYLES.keys())
        raise ValueError(f"未知格式 '{style_name}'，可用格式: {available}")

    cfg = STYLES[style_name]

    # 读取 MD
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    paras = _parse_md(md_text)

    # 创建文档
    doc = Document()

    # 页面设置
    mg = cfg["page_margin_cm"]
    sec = doc.sections[0]
    sec.top_margin = Cm(mg[0])
    sec.bottom_margin = Cm(mg[1])
    sec.left_margin = Cm(mg[2])
    sec.right_margin = Cm(mg[3])

    # 默认样式
    _set_default_style(doc, cfg["body"])

    # 构建内容
    in_code = False
    code_lines = []

    for para in paras:
        t = para['type']

        if t == 'spacer':
            continue

        if t == 'codeblock':
            if '```' in para.get('text', '') and not in_code:
                in_code = True
                code_lines = []
            elif '```' in para.get('text', '') and in_code:
                in_code = False
                cfg_code = cfg["code"]
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.line_spacing = cfg_code["line_spacing"]
                    _add_run(p, cl, cfg_code["font"], cfg_code["size_pt"])
                code_lines = []
            continue

        if in_code:
            code_lines.append(para.get('text', ''))
            continue

        if t == 'hr':
            continue

        if t == 'heading':
            level = para['level']
            text = para.get('text', '')
            cfg_h = cfg["heading"].get(level, cfg["heading"][3])

            p = doc.add_paragraph()
            p.alignment = _align_from_str(cfg_h.get("align", "left"))
            p.paragraph_format.space_before = Pt(cfg_h.get("space_before", 6))
            p.paragraph_format.space_after = Pt(cfg_h.get("space_after", 3))
            p.paragraph_format.line_spacing = cfg["body"]["line_spacing"]

            _add_run(p, text, cfg_h["font"], cfg_h["size_pt"],
                     bold=cfg_h.get("bold", True))
            continue

        if t == 'bullet':
            text = para.get('text', '')
            cfg_b = cfg["bullet"]
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = cfg_b["line_spacing"]
            _add_run(p, '• ' + text if not text.startswith('•') else text,
                     cfg_b["font"], cfg_b["size_pt"])
            continue

        if t == 'quote':
            text = para.get('text', '')
            cfg_q = cfg["quote"]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(cfg_q.get("left_indent_cm", 1.0))
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = cfg_q["line_spacing"]
            _add_run(p, text, cfg_q["font"], cfg_q["size_pt"])
            continue

        if t == 'table':
            rows_data = para.get('rows', [])
            if len(rows_data) >= 2:
                # 解析表头（第一行）
                header_cells = [c.strip() for c in rows_data[0].split('|')[1:-1]]
                # 解析数据行（跳过表头后的分隔行，即含 --- 的行）
                data_rows_raw = [l for l in rows_data[1:] if not re.match(r'^[\s\-:|]+\|?$', l)]
                data_cells = []
                for dr in data_rows_raw:
                    cells = [c.strip() for c in dr.split('|')[1:-1]]
                    if cells:
                        data_cells.append(cells)

                if header_cells and data_cells:
                    from docx.shared import Pt as TablePt
                    table = doc.add_table(rows=1 + len(data_cells), cols=len(header_cells))
                    table.style = 'Table Grid'
                    # 表头
                    for ci, cell_text in enumerate(header_cells):
                        cell = table.rows[0].cells[ci]
                        cell.text = ''
                        run = cell.paragraphs[0].add_run(cell_text)
                        run.bold = True
                        run.font.name = cfg_body["font"]
                        run.font.size = Pt(cfg_body["size_pt"])
                        cell.paragraphs[0].alignment = _align_from_str("center")
                    # 数据行
                    for ri, row_cells in enumerate(data_cells):
                        for ci, cell_text in enumerate(row_cells):
                            cell = table.rows[ri + 1].cells[ci]
                            cell.text = ''
                            run = cell.paragraphs[0].add_run(cell_text)
                            run.font.name = cfg_body["font"]
                            run.font.size = Pt(cfg_body["size_pt"])
                            # 第一列居中，其余左对齐
                            if ci == 0:
                                cell.paragraphs[0].alignment = _align_from_str("center")
                            else:
                                cell.paragraphs[0].alignment = _align_from_str("left")
                    # 表后空行
                    doc.add_paragraph()
            continue

        # paragraph — core body
        if t == 'paragraph':
            text = para.get('text', '')
            cfg_body = cfg["body"]
            p = doc.add_paragraph()
            if cfg.get("first_line_indent", True):
                p.paragraph_format.first_line_indent = Pt(cfg_body["indent_pt"])
            p.paragraph_format.line_spacing = cfg_body["line_spacing"]
            p.alignment = _align_from_str("justify")

            # 处理加粗
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    bold_text = part[2:-2]
                    cfg_bold = cfg.get("bold", cfg_body)
                    _add_run(p, bold_text, cfg_bold["font"],
                             cfg_bold.get("size_pt", cfg_body["size_pt"]),
                             bold=True)
                else:
                    _add_run(p, part, cfg_body["font"], cfg_body["size_pt"])

    # 保存
    doc.save(output_path)

    # 估算页数
    approx_chars = len(md_text)
    chars_per_page = 1800 if cfg_body["size_pt"] >= 12 else 2500
    approx_pages = max(1, round(approx_chars / chars_per_page))

    return {
        "status": "ok",
        "output": output_path,
        "style": style_name,
        "style_label": cfg["label"],
        "pages_approx": approx_pages,
        "chars_total": approx_chars,
    }


# ============================================================
# CLI 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="中文学术论文 .docx 格式转换器 — create_formatted_docx.py")
    parser.add_argument("action", nargs="?", default="convert",
                        choices=["convert", "styles"])
    parser.add_argument("--input", "-i", help="输入 .md 文件路径")
    parser.add_argument("--output", "-o", help="输出 .docx 文件路径")
    parser.add_argument("--style", "-s", default="default",
                        help=f"格式模板（可用: {', '.join(STYLES.keys())}）")

    # 也支持 positional: convert <input> <output> --style X
    parser.add_argument("pos_input", nargs="?")
    parser.add_argument("pos_output", nargs="?")

    args = parser.parse_args()

    # styles 命令
    if args.action == "styles":
        print("=" * 60)
        print("可用格式模板:")
        print("=" * 60)
        for key, cfg in STYLES.items():
            print(f"\n  {key:12s} — {cfg['label']}")
            print(f"              {cfg['desc']}")
            bg = cfg["body"]
            print(f"              正文: {bg['font']} {bg['size_pt']}pt, "
                  f"行距{cfg['body']['line_spacing']}")
            h1 = cfg["heading"][1]
            print(f"              一级标题: {h1['font']} {h1['size_pt']}pt, "
                  f"{h1.get('align','left')}")
        print()
        return

    # convert 命令
    input_path = args.input or args.pos_input
    output_path = args.output or args.pos_output

    if not input_path or not output_path:
        # 默认值
        input_path = args.input or r"d:\公共管理科研\10_研究输出\定稿\技术在场服务缺席_养老服务数字排斥与包容性治理_终稿.md"
        output_path = args.output or r"d:\公共管理科研\10_研究输出\定稿\技术在场服务缺席_养老服务数字排斥与包容性治理_终稿.docx"

    try:
        result = md_to_docx(input_path, output_path, args.style)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)},
                         ensure_ascii=False), file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
